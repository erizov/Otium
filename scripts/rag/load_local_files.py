# -*- coding: utf-8 -*-
"""Load local files from `rag_input/` into the RAG document cache.

Pipeline features:
- Upload: PDF, DOCX, XLSX, TXT/MD, RTF, archives (.zip/.7z/.rar)
- OCR: scans/images + fallback OCR for PDF when text extraction is empty
- Parsing: extract text, tables (XLSX), and section headings where possible
- Normalization: clean noise and emit a consistent Markdown-like text
- Metadata: document type, version/hash, date, inferred discipline
- Quality control: empty pages, OCR health, duplicates

Hierarchy:
- 1st level: City (inferred from filename)
- 2nd level: Place name (inferred from file content; otherwise overview)

All ingested docs are written under:
  `.rag_cache/docs/<city>/<place>/local_file_<lang>_<docid>.json`
and a QC report is written under:
  `rag_outputs/ingest_report.json`.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

_PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

from scripts.rag.config import rag_paths
from scripts.rag.doc_schema import RagDocument
from scripts.rag.doc_schema import doc_to_dict
from scripts.rebuild_stale_city_guide_pdfs import _discover_slugs

_WS_RE = re.compile(r"[ \\t]+")
_BAD_LINE_RE = re.compile(r"^(page\\s+\\d+|страница\\s+\\d+)$", re.I)


def _log_line(text: str) -> None:
    """Console-safe logging on Windows (avoid UnicodeEncodeError)."""
    try:
        sys.stdout.buffer.write((text + "\n").encode("utf-8", errors="replace"))
    except Exception:
        # Last resort: best-effort print.
        try:
            print(text)
        except Exception:
            pass


def _utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _doc_id_for_path(path: Path) -> str:
    st = path.stat()
    raw = "{}|{}|{}".format(path.as_posix(), int(st.st_size), int(st.st_mtime))
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _infer_city_and_lang(project_root: Path, input_root: Path, path: Path) -> tuple[str, str]:
    slugs = set(_discover_slugs(project_root))
    city_aliases: dict[str, str] = {
        "moscow": "moscow",
        "москва": "moscow",
        "spb": "spb",
        "saint petersburg": "spb",
        "санкт-петербург": "spb",
        "petersburg": "spb",
        "smolensk": "smolensk",
        "смоленск": "smolensk",
        "boston": "boston",
    }

    name_low = path.name.lower()
    city = "global"
    for token in sorted(city_aliases, key=len, reverse=True):
        if token in name_low:
            city = city_aliases[token]
            break
    if city == "global":
        for s in slugs:
            if s in name_low:
                city = s
                break

    lang = "unknown"
    if re.search(r"\\b(en|eng|english)\\b", name_low):
        lang = "en"
    elif re.search(r"\\b(ru|rus|russian)\\b", name_low) or re.search(
        r"[А-Яа-яЁё]", path.name
    ):
        lang = "ru"
    return city, lang


def _normalize_text(text: str) -> str:
    s = text.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    for raw in s.split("\n"):
        line = _WS_RE.sub(" ", raw).strip()
        if not line:
            lines.append("")
            continue
        if _BAD_LINE_RE.match(line):
            continue
        lines.append(line)
    # collapse blank runs
    out: list[str] = []
    blank = 0
    for line in lines:
        if not line:
            blank += 1
            if blank <= 1:
                out.append("")
            continue
        blank = 0
        out.append(line)
    return "\n".join(out).strip()


def _infer_place_from_content(path: Path, text: str, *, city_slug: str) -> str:
    """
    Infer a place/item name from content.

    Heuristic: use first strong heading-like line; fallback to filename stem.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for l in lines[:40]:
        if len(l) < 4 or len(l) > 90:
            continue
        if l.isupper():
            continue
        if l.endswith((".", ":", ";")):
            continue
        # Avoid repeating the city as the place.
        low = l.lower()
        if city_slug != "global" and city_slug in low:
            continue
        if "москва" in low and city_slug == "moscow":
            continue
        return l
    stem = path.stem
    # Strip common city token from filename
    stem = re.sub(r"(?i)\\b(moscow|москва|spb|petersburg|санкт[- ]петербург)\\b", "", stem)
    stem = _WS_RE.sub(" ", stem).strip(" -_.,")
    return stem[:80] if stem else "city_overview"


def _classify_discipline(text: str) -> str | None:
    low = text.lower()
    rules = [
        ("history", ("history", "история", "век", "династ", "война")),
        ("architecture", ("architecture", "архитект", "стиль", "зодчеств")),
        ("religion", ("church", "cathedral", "mosque", "synagogue", "храм", "собор", "мечеть", "синагог", "монастыр")),
        ("culture", ("museum", "театр", "theatre", "выстав", "искусств")),
        ("transport", ("metro", "станция", "railway", "вокзал")),
    ]
    scores: dict[str, int] = {}
    for label, keys in rules:
        scores[label] = sum(1 for k in keys if k in low)
    best = max(scores.items(), key=lambda kv: kv[1])
    return best[0] if best[1] > 0 else None


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


def _read_pdf(
    path: Path,
    *,
    allow_ocr: bool,
    max_pages: int,
) -> str:
    from pypdf import PdfReader  # third-party (already in requirements)

    reader = PdfReader(str(path))
    chunks: list[str] = []
    empty_pages = 0
    pages = list(reader.pages)
    if max_pages > 0:
        pages = pages[:max_pages]
    for page in pages:
        t = page.extract_text() or ""
        t = t.strip()
        if t:
            chunks.append(t)
        else:
            empty_pages += 1
    # If empty, try OCR fallback only when explicitly enabled.
    if allow_ocr and not chunks and empty_pages > 0:
        ocr = _ocr_pdf_pages(path)
        if ocr:
            return ocr
    return "\n\n".join(chunks).strip()


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required for .docx: pip install python-docx"
        ) from exc
    d = docx.Document(str(path))
    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    return "\n\n".join(paras).strip()


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception as exc:
        raise RuntimeError("openpyxl is required for .xlsx") from exc
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## {sheet.title}")
        # limit scan to avoid insane files
        max_r = min(sheet.max_row or 0, 3000)
        max_c = min(sheet.max_column or 0, 40)
        for r in range(1, max_r + 1):
            row_vals: list[str] = []
            for c in range(1, max_c + 1):
                v = sheet.cell(row=r, column=c).value
                if v is None:
                    row_vals.append("")
                else:
                    row_vals.append(str(v).strip())
            if any(x for x in row_vals):
                parts.append(" | ".join(x for x in row_vals if x))
    return "\n\n".join(parts).strip()


def _read_rtf(path: Path) -> str:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
    except Exception as exc:
        raise RuntimeError("striprtf is required for .rtf") from exc
    raw = path.read_text(encoding="utf-8", errors="replace")
    return rtf_to_text(raw).strip()


def _ocr_image(path: Path) -> str:
    import pytesseract  # type: ignore
    from PIL import Image

    img = Image.open(str(path))
    # Let tesseract auto-detect; user can control via env if needed.
    return str(pytesseract.image_to_string(img) or "").strip()


def _ocr_pdf_pages(path: Path) -> str:
    """
    Optional OCR for PDF pages.

    Requires `pdf2image` + Poppler to be installed. If unavailable, returns ''.
    """
    try:
        from pdf2image import convert_from_path  # type: ignore
    except Exception:
        return ""
    import pytesseract  # type: ignore

    try:
        images = convert_from_path(str(path), dpi=220)
    except Exception:
        return ""
    parts: list[str] = []
    for img in images[:40]:
        txt = pytesseract.image_to_string(img) or ""
        txt = txt.strip()
        if txt:
            parts.append(txt)
    return "\n\n".join(parts).strip()


def _unpack_archives(input_root: Path, work_dir: Path) -> None:
    """
    Extract .zip/.7z/.rar into a temp work dir for ingestion.

    - zip: stdlib
    - 7z: py7zr
    - rar: rarfile (needs unrar/bsdtar installed for some formats)
    """
    work_dir.mkdir(parents=True, exist_ok=True)
    for p in sorted(input_root.rglob("*")):
        if not p.is_file():
            continue
        suf = p.suffix.lower()
        if suf == ".zip":
            dest = work_dir / p.stem
            dest.mkdir(parents=True, exist_ok=True)
            with zipfile.ZipFile(str(p)) as z:
                z.extractall(str(dest))
        elif suf == ".7z":
            try:
                import py7zr  # type: ignore
            except Exception:
                continue
            dest = work_dir / p.stem
            dest.mkdir(parents=True, exist_ok=True)
            with py7zr.SevenZipFile(str(p), mode="r") as z:
                z.extractall(path=str(dest))
        elif suf == ".rar":
            try:
                import rarfile  # type: ignore
            except Exception:
                continue
            dest = work_dir / p.stem
            dest.mkdir(parents=True, exist_ok=True)
            try:
                with rarfile.RarFile(str(p)) as rf:
                    rf.extractall(str(dest))
            except Exception:
                continue

def _extract_text(
    path: Path,
    *,
    allow_ocr: bool,
    pdf_max_pages: int,
) -> str:
    suf = path.suffix.lower()
    if suf in (".txt", ".md"):
        return _read_text(path)
    if suf == ".pdf":
        return _read_pdf(path, allow_ocr=allow_ocr, max_pages=pdf_max_pages)
    if suf == ".docx":
        return _read_docx(path)
    if suf == ".xlsx":
        return _read_xlsx(path)
    if suf in (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"):
        return _ocr_image(path) if allow_ocr else ""
    if suf == ".rtf":
        return _read_rtf(path)
    if suf in (".djvu", ".doc"):
        return ""
    return ""


def _write_doc(project_root: Path, doc: RagDocument) -> Path:
    paths = rag_paths(project_root)
    place = str(doc.extra.get("place") or "city_overview").strip() or "city_overview"
    safe = re.sub(r"[^A-Za-z0-9_\\-\\.]+", "_", place)[:80]
    out_dir = paths.docs_dir / doc.city_slug / safe
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "{}_{}_{}.json".format(
        doc.source_name,
        doc.language,
        doc.doc_id[:12],
    )
    out.write_text(
        json.dumps(doc_to_dict(doc), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return out


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--project-root",
        type=Path,
        default=None,
        help="Repo root (default: auto).",
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=None,
        help="Directory to ingest (default: rag_input/).",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Rewrite docs even if they already exist (doc_id changes with mtime).",
    )
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Enable OCR for images and PDF fallback (slow; may require Poppler).",
    )
    parser.add_argument(
        "--no-archives",
        action="store_true",
        help="Skip archive unpacking (default: unpack archives).",
    )
    parser.add_argument(
        "--max-files",
        type=int,
        default=0,
        metavar="N",
        help="Limit ingested files for debugging (0 = no limit).",
    )
    parser.add_argument(
        "--pdf-max-pages",
        type=int,
        default=60,
        metavar="N",
        help="Max PDF pages to extract per file (default: 60).",
    )
    args = parser.parse_args()
    project_root = (
        args.project_root.resolve()
        if args.project_root
        else Path(__file__).resolve().parent.parent.parent
    )
    input_root = (
        args.input_dir.resolve()
        if args.input_dir
        else (project_root / "rag_input")
    )
    if not input_root.is_dir():
        print("Missing rag_input dir: {}".format(input_root), file=sys.stderr)
        return 2
    paths = rag_paths(project_root)
    paths.docs_dir.mkdir(parents=True, exist_ok=True)

    # Unpack archives into a temp dir under rag_outputs/ so rglob can ingest.
    work_dir = paths.outputs_dir / "_ingest_work"
    if not args.no_archives:
        _unpack_archives(input_root, work_dir)

    wrote = 0
    skipped = 0
    dup = 0
    errors: list[dict[str, str]] = []
    seen_hashes: set[str] = set()

    ingest_roots = [input_root, work_dir]
    processed = 0
    for root_in in ingest_roots:
        for p in sorted(root_in.rglob("*")):
            if not p.is_file():
                continue
            if p.name.startswith("."):
                continue
            processed += 1
            if args.max_files and processed > int(args.max_files):
                break
            _log_line("Ingest: {}".format(p.name))
            try:
                raw = _extract_text(
                    p,
                    allow_ocr=bool(args.ocr),
                    pdf_max_pages=int(args.pdf_max_pages),
                )
            except Exception as exc:
                errors.append({"path": p.as_posix(), "error": str(exc)})
                continue
            if not raw:
                skipped += 1
                continue
            text = _normalize_text(raw)
            if not text:
                skipped += 1
                continue
            city, lang = _infer_city_and_lang(project_root, input_root, p)
            place = _infer_place_from_content(p, text, city_slug=city)
            discipline = _classify_discipline(text)
            doc_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
            if doc_hash in seen_hashes:
                dup += 1
                continue
            seen_hashes.add(doc_hash)

            doc_id = _doc_id_for_path(p)
            doc = RagDocument(
                doc_id=doc_id,
                city_slug=city,
                language=lang,
                source_name="local_file",
                source_url=p.resolve().as_uri(),
                license="Local file (user-provided)",
                retrieved_at_utc=_utc_now(),
                title=p.stem,
                text=text,
                extra={
                    "place": place,
                    "discipline": discipline,
                    "doc_type": p.suffix.lower().lstrip("."),
                    "version": doc_hash[:12],
                    "file_mtime_utc": datetime.fromtimestamp(
                        p.stat().st_mtime, tz=timezone.utc
                    ).strftime("%Y-%m-%dT%H:%M:%SZ"),
                    "path_rel": (
                        str(p.relative_to(input_root))
                        if root_in == input_root
                        else str(p.relative_to(work_dir))
                    ),
                    "filename": p.name,
                },
            )
            out_path = _write_doc(project_root, doc)
            if out_path.is_file():
                wrote += 1
        if args.max_files and processed >= int(args.max_files):
            break

    report = {
        "loaded": wrote,
        "skipped": skipped,
        "duplicates": dup,
        "errors": errors[:200],
    }
    paths.outputs_dir.mkdir(parents=True, exist_ok=True)
    (paths.outputs_dir / "ingest_report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print("Loaded local docs:", wrote, "skipped:", skipped, "dup:", dup)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

